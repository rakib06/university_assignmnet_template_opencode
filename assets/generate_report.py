#!/usr/bin/env python3
"""
Reusable Academic DOCX Report Generator
========================================
Generates a professionally formatted university assignment report.

Usage:
    python3 assets/generate_report.py

Reads configuration from: templates/report_config.json
Reads source code from:    *.asm, *.py, *.c, *.java (auto-detected)
Embeds screenshots from:   screenshots/ folder
Outputs:                   report_final.docx
"""

import glob
import json
import os
import sys
from datetime import datetime

try:
    from PIL import Image
except ImportError:
    Image = None
    print("[WARN] Pillow not installed. Image sizing may be approximate.")
    print("       Install with: pip install Pillow")

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ══════════════════════════════════════════════════════════════════════
#  PATHS
# ══════════════════════════════════════════════════════════════════════

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CONFIG_FILE = os.path.join(BASE_DIR, "templates", "report_config.json")
SCREENSHOTS_DIR = os.path.join(BASE_DIR, "screenshots")
ASSETS_DIR = os.path.join(BASE_DIR, "assets")
LOGO_FILE = os.path.join(ASSETS_DIR, "bup_logo.png")
OUTPUT_FILE = os.path.join(BASE_DIR, "report_final.docx")
PAGE_WIDTH_INCHES = 6.5


# ══════════════════════════════════════════════════════════════════════
#  CONFIGURATION LOADER
# ══════════════════════════════════════════════════════════════════════

def load_config():
    """Load assignment configuration from report_config.json."""
    if not os.path.exists(CONFIG_FILE):
        print(f"[ERROR] Configuration file not found: {CONFIG_FILE}")
        print("        Run setup.py first to generate the configuration.")
        sys.exit(1)
    with open(CONFIG_FILE, "r") as f:
        return json.load(f)


def find_source_code():
    """Auto-detect the main source code file in the project root."""
    extensions = ["*.asm", "*.py", "*.c", "*.cpp", "*.java", "*.js", "*.go"]
    for ext in extensions:
        files = glob.glob(os.path.join(BASE_DIR, ext))
        if files:
            return files[0]
    return None


# ══════════════════════════════════════════════════════════════════════
#  HELPERS
# ══════════════════════════════════════════════════════════════════════

def setup_styles(doc):
    """Configure document-wide styles: Times New Roman, black, 1.5 spacing."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.name = "Times New Roman"
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.font.bold = True
        hs.font.size = Pt(16 if i == 1 else 14 if i == 2 else 12)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

    try:
        lb = doc.styles["List Bullet"]
        lb.font.name = "Times New Roman"
        lb.font.size = Pt(12)
        lb.font.color.rgb = RGBColor(0, 0, 0)
    except KeyError:
        pass


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = "Times New Roman"
    return h


def add_body(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic
    return p


def add_code_block(doc, code_text, font_size=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def get_image_dimensions(filepath):
    """Return (width_inches, height_inches) scaled to fit page."""
    if Image:
        with Image.open(filepath) as img:
            w_px, h_px = img.size
        aspect = h_px / w_px
        w = min(w_px / 96.0, PAGE_WIDTH_INCHES)
        h = w * aspect
        if h > 7.5:
            h = 7.5
            w = h / aspect
        return w, h
    else:
        return PAGE_WIDTH_INCHES, 4.0


def try_add_image(doc, filename, caption):
    """Embed an image with proper sizing, or add a grey placeholder."""
    path = os.path.join(SCREENSHOTS_DIR, filename)
    if not os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[ Screenshot: {filename} ]")
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        return False

    w, h = get_image_dimensions(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(w))

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.font.size = Pt(10)
        r.font.italic = True
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)
    return True


def add_table(doc, headers, rows, header_color="D9E2F3"):
    """Add a styled table with Times New Roman text."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{header_color}"/>')
        cell._element.get_or_add_tcPr().append(shading)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("")
    return table


# ══════════════════════════════════════════════════════════════════════
#  SECTION BUILDERS
# ══════════════════════════════════════════════════════════════════════

def build_cover_page(doc, config):
    """Build the assignment cover page."""
    for _ in range(3):
        doc.add_paragraph("")

    if os.path.exists(LOGO_FILE):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            w, _ = get_image_dimensions(LOGO_FILE)
            run = p.add_run()
            run.add_picture(LOGO_FILE, width=Inches(min(w, 2.0)))
        except Exception:
            p.add_run("[ University Logo ]")
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(config.get("assignment_title", "Assignment Report"))
    run.font.size = Pt(24)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("")

    course_line = f"{config.get('course_code', '')} - {config.get('course_name', '')}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(course_line)
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(config.get("university", ""))
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("")

    for line in [
        "Submitted by:",
        "",
        f"Name : {config.get('student_name', '')}",
        f"ID   : {config.get('student_id', '')}",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)

    if config.get("semester") or config.get("section"):
        doc.add_paragraph("")
        sem_sec = f"{config.get('semester', '')}"
        if config.get("section"):
            sem_sec += f"  |  Section: {config['section']}"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(sem_sec)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)

    if config.get("faculty_name"):
        doc.add_paragraph("")
        fac = f"Submitted to: {config['faculty_name']}"
        if config.get("faculty_designation"):
            fac += f", {config['faculty_designation']}"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(fac)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)

    if config.get("submission_date"):
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Date: {config['submission_date']}")
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_page_break()


def build_toc(doc):
    """Build Table of Contents."""
    add_heading(doc, "Table of Contents", level=1)
    items = [
        "1. Introduction",
        "2. Program Logic (Algorithm)",
        "3. Source Code Listing",
        "4. Detailed Code Explanation by Segment",
        "5. Instructions Reference",
        "6. Screenshots",
        "7. Expected Textual Output",
        "8. Conclusion",
        "References",
    ]
    for item in items:
        p = doc.add_paragraph(item)
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_page_break()


def build_screenshots_section(doc):
    """Build the screenshots section by auto-detecting available images."""
    add_heading(doc, "6. Screenshots", level=1)
    add_body(doc,
        "This section presents visual evidence of the program's compilation, "
        "execution, and output as captured from the development environment."
    )

    screenshot_map = [
        ("code_view.png", "6.1", "Source Code in Editor (Top Half)",
         "The top portion of the source code as displayed in the editor window."),
        ("code_view_2.png", "6.2", "Source Code in Editor (Bottom Half)",
         "The lower portion of the source code after scrolling down."),
        ("compile.png", "6.3", "Compilation Result",
         "The assembler output after compilation, showing any warnings or errors."),
        ("run.png", "6.4", "Execution — Run Button",
         "The run button highlighted after successful compilation."),
        ("run_start.png", "6.5", "Console — Input Prompt",
         "The console window displaying the input prompt."),
        ("execution_output.png", "6.6", "Console — Full Output",
         "The complete execution output showing results."),
        ("error1.png", "6.7", "Error Encountered",
         "Any error encountered during compilation or execution."),
    ]

    fig_num = 1
    for filename, sec, title, desc in screenshot_map:
        filepath = os.path.join(SCREENSHOTS_DIR, filename)
        if os.path.exists(filepath):
            add_heading(doc, f"{sec} {title}", level=2)
            add_body(doc, desc)
            try_add_image(doc, filename, f"Figure {fig_num}: {title}")
            fig_num += 1


# ══════════════════════════════════════════════════════════════════════
#  MAIN REPORT BUILDER
# ══════════════════════════════════════════════════════════════════════

def build_report():
    """Build the complete assignment report."""
    config = load_config()
    doc = Document()
    setup_styles(doc)

    # ── Cover Page ───────────────────────────────────────────────────
    build_cover_page(doc, config)

    # ── Table of Contents ────────────────────────────────────────────
    build_toc(doc)

    # ── 1. Introduction ──────────────────────────────────────────────
    add_heading(doc, "1. Introduction", level=1)
    add_body(doc,
        f"This report presents the implementation of the assignment titled "
        f"\"{config.get('assignment_title', 'N/A')}\" as part of the course "
        f"{config.get('course_code', '')} - {config.get('course_name', '')} "
        f"at {config.get('university', '')}."
    )
    add_body(doc,
        "The objective of this assignment is to develop, test, and document a "
        "program using the specified platform. The report provides a comprehensive "
        "explanation of the program's design, logic, implementation details, and "
        "execution results."
    )
    add_body(doc,
        "The following sections describe the algorithmic approach, present the "
        "complete source code, explain each code segment in detail, reference "
        "all instructions used, and demonstrate the program's correct operation "
        "through execution screenshots."
    )

    # ── 2. Program Logic ─────────────────────────────────────────────
    add_heading(doc, "2. Program Logic (Algorithm)", level=1)
    add_body(doc,
        "The program logic is decomposed into the following sequential steps. "
        "Each step maps to a specific block of assembly instructions."
    )
    steps = [
        ("Step 1 — Displaying the Input Prompt",
         "The program displays a prompt message to the console using DOS "
         "interrupt INT 21h with function AH = 09h, which prints a dollar-terminated string."),
        ("Step 2 — Reading User Input",
         "A single keystroke is read using INT 21h function AH = 01h. The ASCII "
         "character is converted to a numeric value by subtracting '0' (ASCII 48)."),
        ("Step 3 — Initializing Registers",
         "The program initializes the necessary registers for computation: "
         "the accumulator, base register, sum register, and loop counter."),
        ("Step 4 — Computation Loop",
         "The core computation loop executes N times. On each iteration, the "
         "program performs the required calculation, saves registers, displays "
         "the current result, adds separators, restores registers, and computes "
         "the next value."),
        ("Step 5 — Displaying Results",
         "After the loop completes, the program displays the final result "
         "(e.g., sum) and formats the output appropriately."),
        ("Step 6 — Program Termination",
         "The program invokes INT 21h function AH = 4Ch to terminate cleanly."),
    ]
    for title_text, detail in steps:
        p = doc.add_paragraph()
        run = p.add_run(title_text + "\n")
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        detail_run = p.add_run(detail)
        detail_run.font.name = "Times New Roman"
        detail_run.font.size = Pt(12)
        detail_run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_page_break()

    # ── 3. Source Code Listing ───────────────────────────────────────
    add_heading(doc, "3. Source Code Listing", level=1)
    add_body(doc,
        "The complete source code is presented below. The program is organized "
        "into a data segment for variable declarations and a code segment for "
        "executable instructions."
    )
    source_file = find_source_code()
    if source_file:
        with open(source_file, "r") as f:
            code = f.read()
        add_code_block(doc, code, font_size=8)
    else:
        add_body(doc, "[ No source code file found in the project root ]")

    doc.add_page_break()

    # ── 4. Detailed Code Explanation ─────────────────────────────────
    add_heading(doc, "4. Detailed Code Explanation by Segment", level=1)
    add_body(doc,
        "This section provides a comprehensive, segment-by-segment explanation "
        "of the source code. Each subsection presents the relevant code snippet "
        "followed by a detailed discussion of its function, the rationale behind "
        "the approach, and the mechanism by which the processor executes it."
    )
    add_body(doc,
        "[ The AI agent will generate detailed segment explanations here when "
        "you say \"Explain each code segment in academic tone\" ]"
    )

    doc.add_page_break()

    # ── 5. Instructions Reference ────────────────────────────────────
    add_heading(doc, "5. Instructions Reference", level=1)
    add_body(doc,
        "The following table lists every instruction and directive used in "
        "this program with a description and an example."
    )
    add_body(doc,
        "[ The AI agent will populate this table based on your specific code ]"
    )

    doc.add_page_break()

    # ── 6. Screenshots ───────────────────────────────────────────────
    build_screenshots_section(doc)

    doc.add_page_break()

    # ── 7. Expected Output ───────────────────────────────────────────
    add_heading(doc, "7. Expected Textual Output", level=1)
    add_body(doc,
        "The following is the expected textual output when the program is "
        "executed with appropriate input."
    )
    add_body(doc,
        "[ The AI agent will add the sample output here after you run the program ]"
    )

    # ── 8. Conclusion ────────────────────────────────────────────────
    add_heading(doc, "8. Conclusion", level=1)
    add_body(doc,
        "This assignment demonstrates the successful implementation of the "
        "assigned programming task. The key concepts and skills exercised "
        "include:"
    )
    bullets = [
        "Understanding and applying low-level programming constructs.",
        "Implementing loops, conditional branching, and subroutine calls.",
        "Using DOS interrupts for keyboard input and screen output.",
        "Converting between binary and decimal representations.",
        "Designing modular, well-structured code with proper documentation.",
    ]
    for bp in bullets:
        p = doc.add_paragraph(bp, style="List Bullet")
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 0)

    add_body(doc, "")
    add_body(doc,
        "The program was tested and verified to produce correct results. "
        "The execution screenshots provided in Section 6 confirm proper "
        "operation of the program."
    )

    # ── References ───────────────────────────────────────────────────
    add_heading(doc, "References", level=1)
    refs = [
        "Intel Corporation. Intel 8086/8088 User's Manual. 1979.",
        "Microsoft Corporation. MS-DOS Programmer's Reference. 1993.",
        "emu8086 — Microprocessor Emulator. https://emu8086.com",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(f"[{i}] {ref}")
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 0)

    # ── Save ─────────────────────────────────────────────────────────
    doc.save(OUTPUT_FILE)
    print(f"Report saved to: {OUTPUT_FILE}")
    print(f"Configuration: {CONFIG_FILE}")
    print(f"Screenshots: {SCREENSHOTS_DIR}")


if __name__ == "__main__":
    build_report()
